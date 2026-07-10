"""共享 fixtures。"""

import os
import sys
import tempfile
from pathlib import Path

import pytest

# 确保 src/ 在 sys.path 中
_src = Path(__file__).resolve().parent.parent / "src"
if str(_src) not in sys.path:
    sys.path.insert(0, str(_src))


@pytest.fixture
def temp_dir():
    """临时目录，测试结束后自动清理。"""
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


@pytest.fixture
def sample_pdf(temp_dir) -> Path:
    """生成一个最小的示例 PDF 文件。"""
    p = temp_dir / "sample.pdf"
    # 用 pypdf 创建合法 PDF
    from pypdf import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(210, 297)
    writer.write(p)
    return p


@pytest.fixture
def sample_docx(temp_dir) -> Path:
    """生成示例 Word 文档。"""
    p = temp_dir / "sample.docx"
    from docx import Document
    doc = Document()
    doc.add_heading("测试文档", 0)
    doc.add_paragraph("这是第一段测试内容，包含一些技术术语：OTA升级、电源管理。")
    doc.add_paragraph("这是第二段内容。CCU 的诊断功能需要满足 ISO 14229 标准。")
    # 添加表格
    table = doc.add_table(3, 3)
    table.cell(0, 0).text = "需求ID"
    table.cell(0, 1).text = "功能"
    table.cell(0, 2).text = "优先级"
    table.cell(1, 0).text = "REQ-001"
    table.cell(1, 1).text = "OTA升级"
    table.cell(1, 2).text = "高"
    table.cell(2, 0).text = "REQ-002"
    table.cell(2, 1).text = "电源管理"
    table.cell(2, 2).text = "中"
    doc.save(str(p))
    return p


@pytest.fixture
def sample_xlsx(temp_dir) -> Path:
    """生成示例 Excel 文件。"""
    p = temp_dir / "sample.xlsx"
    import pandas as pd
    df = pd.DataFrame({
        "需求ID": ["REQ-001", "REQ-002", "REQ-003"],
        "功能描述": ["OTA升级流程", "电源管理策略", "诊断功能"],
        "优先级": ["高", "中", "低"],
        "状态": ["已实现", "设计中", "规划中"],
    })
    df.to_excel(str(p), index=False, sheet_name="需求矩阵")
    return p


@pytest.fixture
def sample_txt(temp_dir) -> Path:
    """生成示例文本文件。"""
    p = temp_dir / "sample.txt"
    p.write_text(
        "# 系统需求文档\n\n"
        "## 1. OTA 升级\n"
        "车辆应支持通过 4G/5G 网络进行远程固件升级。\n"
        "升级过程中应保证车辆处于 P 档且电池电量 > 50%。\n\n"
        "## 2. 电源管理\n"
        "CCU 应支持三种电源模式：睡眠、待机、运行。\n"
        "睡眠模式下功耗应低于 1mA。\n\n"
        "## 3. 诊断功能\n"
        "系统应支持 UDS (ISO 14229) 标准诊断服务。\n"
        "包括读取/清除 DTC、读取数据流、执行例程等。\n",
        encoding="utf-8",
    )
    return p


@pytest.fixture
def sample_md(temp_dir) -> Path:
    """生成示例 Markdown 文件。"""
    p = temp_dir / "sample.md"
    p.write_text(
        "# 项目需求\n\n"
        "## 功能需求\n\n"
        "- **FR-01**: OTA 远程升级，支持断点续传\n"
        "- **FR-02**: 电源管理，支持低功耗模式\n"
        "- **FR-03**: CAN 总线通信，500kbps\n\n"
        "## 非功能需求\n\n"
        "- **NFR-01**: 启动时间 < 3 秒\n"
        "- **NFR-02**: 静态功耗 < 1mA\n",
        encoding="utf-8",
    )
    return p


@pytest.fixture
def sample_project(temp_dir) -> Path:
    """生成包含多种文档的项目目录。"""
    project = temp_dir / "project"
    project.mkdir()
    # TXT
    (project / "readme.txt").write_text(
        "OTA 升级流程设计文档 v2.0\n\n"
        "升级包下载完成后进行 CRC 校验，校验通过后写入备份分区。\n"
        "写入完成后设置启动标志位，下次上电从新分区启动。\n"
        "如果新分区启动失败，自动回滚到备份分区。\n",
        encoding="utf-8",
    )
    # MD
    (project / "design.md").write_text(
        "# 电源管理设计\n\n"
        "## 状态机\n"
        "睡眠 → 待机 → 运行 → 待机 → 睡眠\n\n"
        "## 功耗指标\n"
        "| 模式 | 功耗 |\n"
        "|------|------|\n"
        "| 睡眠 | <1mA |\n"
        "| 待机 | <50mA |\n"
        "| 运行 | <500mA |\n",
        encoding="utf-8",
    )
    return project


def pytest_configure(config):
    config.addinivalue_line("markers", "slow: 慢速测试")
    config.addinivalue_line("markers", "network: 需要网络/Ollama 服务")


# ---- Ollama 连通性检查 ----
_ollama_url = os.getenv("EMBED_BASE_URL", "http://localhost:11434")
_embed_model = os.getenv("EMBED_MODEL_NAME", "bge-m3")


def _check_ollama():
    try:
        import httpx
        r = httpx.get(_ollama_url, timeout=3.0)
        if r.status_code != 200:
            return False
        r = httpx.post(
            f"{_ollama_url}/api/embeddings",
            json={"model": _embed_model, "prompt": "test"},
            timeout=5.0,
        )
        return r.status_code == 200
    except Exception:
        return False


require_ollama = pytest.mark.skipif(
    not _check_ollama(),
    reason=f"Ollama Embedding 不可用 ({_ollama_url}, model={_embed_model})",
)
